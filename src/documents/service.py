from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sqlalchemy.orm import Session
from typing import Tuple, Dict, List
import os
import time

from src.documents.models import Document
from src.vector_store.client import vector_store
from src.core.config import get_settings
from src.core.logging import logger, log_performance


settings = get_settings()


class DocumentService:
    """Service for document operations"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap= settings.CHUNK_OVERLAP,
            length_function = len,
            separators=["\n\n","\n"," ",""]
        )
        
    def upload_pdf(self, file_path: str, title:str, description: str, db:Session):
        """Upload and process a PDF document"""
        start_time = time.time()
        
        try:
            text, file_metadata = self.extract_text(file_path)
            
            document = Document(
                title=title,
                description= description,
                file_name = os.path.basename(file_path),
                file_path = file_path,
                file_size = file_metadata['file_size'],
                page_count = file_metadata['page_count']
            )
            
            db.add(document)
            db.flush()
            
            chunks = self.chunk_text(text)
            
            chunk_ids = self.store_embeddings(
                document_id = document.id,
                title = title,
                chunks = chunks
            )
            
            document.chunk_count = len(chunk_ids)
            db.commit()
            db.refresh(document)
            
            processing_time = time.time() - start_time
            stats = {
                "document_id" : document.id,
                "chunks_created" : len(chunk_ids),
                "text_length" : len(text),
                "processing_time" : processing_time
            }
            
            log_performance("upload_pdf", processing_time, document_id = document.id)
            logger.info(f"Document uploaded: {document.id} - {title}")
            
            return document, stats
        
        except Exception as e:
            db.rollback()
            logger.error(f"Error uploading PDF: {str(e)}")
            raise
        
        
    def extract_text(self, file_path: str) -> Tuple[str, Dict]:
        """Extract text from document"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf(file_path)
        
        elif file_extension == '.docx':
            return self._extract_docx(file_path)
        
        elif file_extension == '.txt':
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
    def _extract_pdf(self, file_path:str) -> Tuple[str, Dict]:
        """Extract text from PDF"""
        reader = PdfReader(file_path)
        
        text_parts = []
        for i , page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {i+1}]\n{page_text}")
                
        full_text = "\n\n".join(text_parts)
        
        metadata = {
            "page_count" : len(reader.pages),
            "file_size" : os.path.getsize(file_path)
        }
        
        return full_text, metadata
    
    
    def _extract_docx(self,file_path:str ) -> Tuple[str, Dict]:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        
        text_parts = [para.text for para in doc.paragraphs if para.text]
        full_text = "\n\n".join(text_parts)
        
        metadata = {
            "page_count" : len(doc.sections),
            "file_size" : os.path.getsize(file_path)
        }
        
        return full_text, metadata
    
    
    def _extract_txt(self, file_path:str) -> Tuple[str, Dict]:
        """Extract text from TXT"""
        with open(file_path, 'r', encoding='utf-8') as f:
            full_text = f.read()
            
        metadata = {
            "page_count" : 1,
            "file_size" : os.path.getsize(file_path)
        }
        
        return full_text, metadata
    
    
    def chunk_text(self, text:str ) -> List[str]:
        """Split text into chunks"""
        chunks = self.text_splitter.split_text(text)
        logger.info(f"Text split into {len(chunks)} chunks")
        return chunks
    
    def store_embeddings(self,document_id:str, title:str, chunks: List[str]) -> List[str]:
        """Store embeddings in vector store"""
        metadatas = []
        for i, chunk in enumerate(chunks):
            metadata = {
                "document_id": document_id,
                "title": title,
                "chunk_index" :i,
                "total_chunks" : len(chunks)
            }
            
            metadatas.append(metadata)
            
        chunk_ids = vector_store.add_documents(texts = chunks, metadatas=metadatas)
        
        logger.info(f"Stored {len(chunk_ids)} embeddings for document {document_id}")
        return chunk_ids
    
    
    def delete_document(self, document_id:str , db:Session):
        """Delete document and its embedding"""
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        vector_store.delete_by_document_id(document_id)
        
        if os.path.exists(document.file_path):
            os.remove(document.file_path)
            
        db.delete(document)
        db.commit()
        
        logger.info(f"Deleted document: {document_id}")
        
        
    def get_document(self, document_id: str, db:Session) -> Document:
        """Get document by ID"""
        document = db.query(Document).filter(Document.id == document_id).first()
        
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        return document
    
    def list_documents(self, db:Session , skip:int = 0, limit: int = 100) -> List[Document]:
        """List all documents"""
        return db.query(Document).offset(skip).limit(limit).all()
    
    
document_service = DocumentService()
        